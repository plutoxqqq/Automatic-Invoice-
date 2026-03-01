from __future__ import annotations

import argparse
import calendar
import datetime as dt
import re
import sys
import time
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional

DEFAULT_INVOICE_DIR = Path(
    r"C:\Users\ver0016\OneDrive - Hoppers Crossing Secondary College\Desktop\Study Work\Invoices\New invoices"
    r"C:\Users\ver0016\OneDrive - Hoppers Crossing Secondary College\Desktop\Study Work\Invoices\New invoices"
)


@dataclass(frozen=True)
class InvoiceConfig:
    key: str
    display_name: str
    weekdays: tuple[int, ...]


@dataclass
class ProcessResult:
    source_name: str
    docx_name: str
    pdf_name: str
    invoice_number: str
    subtotal: float
    status: str
    duration_s: float
    error: str = ""


INVOICE_RULES: dict[str, InvoiceConfig] = {
    "AFLO": InvoiceConfig(key="AFLO", display_name="AFLO", weekdays=(calendar.MONDAY,)),
    "Bensons": InvoiceConfig(
        key="Bensons",
        display_name="Bensons",
        weekdays=(calendar.WEDNESDAY, calendar.FRIDAY),
    ),
    "Adeval": InvoiceConfig(key="Adeval", display_name="Adeval", weekdays=(calendar.FRIDAY,)),
    "Rodpak": InvoiceConfig(key="Rodpak", display_name="Rodpak", weekdays=(calendar.SUNDAY,)),
}

# Keep support for spelling variant requested in prior notes.
CUSTOMER_ALIASES = {"Advel": "Adeval"}
# Allow "Advel" typo as requested while still targeting Adeval files.
CUSTOMER_ALIASES = {
    "Advel": "Adeval",
}

REQUIRED_TEMPLATE_FILES = [
    "AFLO Feb.docx",
    "Bensons Feb.docx",
    "Adeval Feb.docx",
    "Rodpak Feb.docx",
]

MONTH_TOKEN = "Feb"


def load_document(docx_path: Path):
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - import guard
        raise SystemExit(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        ) from exc
    return Document(str(docx_path))



def first_weekday_of_month(year: int, month: int, target_weekday: int) -> dt.date:
    for day in range(1, 8):
        candidate = dt.date(year, month, day)
        if candidate.weekday() == target_weekday:
            return candidate
    raise ValueError("Unable to compute first weekday of month")


def all_weekdays_in_month(year: int, month: int, weekdays: Iterable[int]) -> list[dt.date]:
    weekday_set = set(weekdays)
    dates: list[dt.date] = []
    _, days_in_month = calendar.monthrange(year, month)
    for day in range(1, days_in_month + 1):
        current = dt.date(year, month, day)
        if current.weekday() in weekday_set:
            dates.append(current)
    return dates


def add_months(date_value: dt.date, months: int) -> dt.date:
    month = date_value.month - 1 + months
    year = date_value.year + month // 12
    month = month % 12 + 1
    day = min(date_value.day, calendar.monthrange(year, month)[1])
    return dt.date(year, month, day)


def parse_money(value: str) -> float:
    cleaned = value.strip().replace("$", "").replace(",", "")
    if not cleaned:
        return 0.0
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def format_money(value: float) -> str:
    if value.is_integer():
        return f"${int(value):,}"
    return f"${value:,.2f}"


def replace_first_word_with_month(text: str, month_name: str) -> str:
    if not text.strip():
        return text
    parts = text.split(maxsplit=1)
    if len(parts) == 1:
        return month_name
    return f"{month_name} {parts[1]}"


def replace_text_in_runs(paragraph, old: str, new: str) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        return False
    replaced = full_text.replace(old, new)

    cursor = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run.text = replaced[cursor : cursor + run_len]
        cursor += run_len

    if cursor < len(replaced) and paragraph.runs:
        paragraph.runs[-1].text += replaced[cursor:]
    return True


def iter_all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_customer_key(filename: str) -> Optional[str]:
    stem = Path(filename).stem
    for key in INVOICE_RULES:
        if stem.lower().startswith(key.lower()):
            return key
    for alias, real_key in CUSTOMER_ALIASES.items():
        if stem.lower().startswith(alias.lower()):
            return real_key
    return None


def get_due_date(invoice_date: dt.date) -> dt.date:
    return add_months(invoice_date, 1) + dt.timedelta(days=1)


def find_line_with_label(document, label: str):
    lower_label = label.lower()
    for paragraph in iter_all_paragraphs(document):
        if lower_label in paragraph.text.lower():
            return paragraph
    return None


def update_invoice_number(document) -> str:
    paragraph = find_line_with_label(document, "invoice no")
    if not paragraph:
        return ""

    match = re.search(r"([A-Za-z]+)(\d+)", paragraph.text)
    if not match:
        return ""

    prefix, number = match.group(1), match.group(2)
    incremented = f"{prefix}{int(number) + 1:0{len(number)}d}"
    replace_text_in_runs(paragraph, match.group(0), incremented)
    return incremented


def update_labelled_date(document, label: str, new_date: dt.date) -> None:
    paragraph = find_line_with_label(document, label)
    if not paragraph:
        return

    new_text = new_date.strftime("%d/%m/%y")
    new_text = new_date.strftime("%d/%m/%y")

    match = re.search(r"\d{1,2}/\d{1,2}/\d{2}", paragraph.text)
    if match:
        replace_text_in_runs(paragraph, match.group(0), new_text)
        return

    pattern = re.compile(r"(:\s*)([^\s]+)$")
    fallback_match = pattern.search(paragraph.text)
    if fallback_match:
        replace_text_in_runs(paragraph, fallback_match.group(2), new_text)


def update_description(document, month_name: str) -> None:
    paragraph = find_line_with_label(document, "description")
    if not paragraph:
        return

    text = paragraph.text
    idx = text.lower().find("description")
    if idx == -1:
        return

    content_start = text.find(":", idx)
    if content_start == -1:
        return

    original_desc = text[content_start + 1 :].strip()
    content_start += 1
    original_desc = text[content_start:].strip()
    updated_desc = replace_first_word_with_month(original_desc, month_name)
    if original_desc and original_desc != updated_desc:
        replace_text_in_runs(paragraph, original_desc, updated_desc)


def find_service_table(document):
    for table in document.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if len(headers) >= 2 and "date" in headers[0] and "amount" in headers[1]:
            return table
    return None


def duplicate_row(table, row_idx: int):
    row = table.rows[row_idx]
    new_tr = row._tr.clone()  # pylint: disable=protected-access
    table._tbl.append(new_tr)  # pylint: disable=protected-access


def set_service_dates(table, service_dates: list[dt.date]) -> float:
    existing_data_rows = table.rows[1:]

    while len(existing_data_rows) < len(service_dates):
        duplicate_row(table, len(table.rows) - 1)
        existing_data_rows = table.rows[1:]

    while len(existing_data_rows) > len(service_dates):
        table._tbl.remove(existing_data_rows[-1]._tr)  # pylint: disable=protected-access
        existing_data_rows = table.rows[1:]

    subtotal = 0.0
    for row, service_date in zip(existing_data_rows, service_dates):
        if len(row.cells) < 2:
            continue
        row.cells[0].text = service_date.strftime("%d/%m")
        subtotal += parse_money(row.cells[1].text)

    return subtotal


def update_gst_and_total(document, subtotal: float) -> None:
    total_value = subtotal

    for table in document.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                label = cell.text.strip().lower()
                if label == "gst" and idx + 1 < len(row.cells):
                    row.cells[idx + 1].text = "$0"
                if label == "total" and idx + 1 < len(row.cells):
                    row.cells[idx + 1].text = format_money(subtotal)


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return
    except ImportError:
        pass

    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(docx_path))
        document.SaveAs(str(pdf_path), FileFormat=17)
        document.Close()
        word.Quit()
        return
    except Exception as exc:  # pragma: no cover - runtime env specific
        raise RuntimeError(
            "Could not convert DOCX to PDF. Install 'docx2pdf' or ensure Word COM automation works."
        ) from exc


def target_names_for_month(source_path: Path, month_abbrev: str) -> tuple[Path, Path]:
    new_stem = source_path.stem.replace(MONTH_TOKEN, month_abbrev)
    return source_path.with_name(f"{new_stem}.docx"), source_path.with_name(f"{new_stem}.pdf")


def resolve_invoice_files(base_dir: Path) -> list[Path]:
    files: list[Path] = []
    missing: list[str] = []

    for name in REQUIRED_TEMPLATE_FILES:
        candidate = base_dir / name
        if candidate.exists():
            files.append(candidate)
        else:
            missing.append(name)

    if missing:
        raise SystemExit(
            f"Missing required invoice templates in {base_dir}: {', '.join(missing)}."
        )

    return files


def process_invoice(source_doc: Path, year: int, month: int, dry_run: bool = False) -> ProcessResult:
    start = time.perf_counter()
    customer_key = find_customer_key(source_doc.name)
    if not customer_key:
        raise RuntimeError(f"Unknown customer key in filename: {source_doc.name}")

    config = INVOICE_RULES[customer_key]
    invoice_date = first_weekday_of_month(year, month, calendar.SUNDAY)
    due_date = get_due_date(invoice_date)
    service_dates = all_weekdays_in_month(year, month, config.weekdays)

    month_abbrev = dt.date(year, month, 1).strftime("%b")
    month_name = dt.date(year, month, 1).strftime("%B")
    docx_target, pdf_target = target_names_for_month(source_doc, month_abbrev)

    if dry_run:
        return ProcessResult(
            source_name=source_doc.name,
            docx_name=docx_target.name,
            pdf_name=pdf_target.name,
            invoice_number="",
            subtotal=0.0,
            status="planned",
            duration_s=time.perf_counter() - start,
        )

    document = load_document(source_doc)
    update_labelled_date(document, "date", invoice_date)
    update_labelled_date(document, "due date", due_date)
    new_invoice_number = update_invoice_number(document)
    update_description(document, month_name)

    service_table = find_service_table(document)
    subtotal = 0.0
    if service_table:
        subtotal = set_service_dates(service_table, service_dates)
    update_gst_and_total(document, subtotal)

    document.save(str(docx_target))
    convert_to_pdf(docx_target, pdf_target)

    return ProcessResult(
        source_name=source_doc.name,
        docx_name=docx_target.name,
        pdf_name=pdf_target.name,
        invoice_number=new_invoice_number,
        subtotal=subtotal,
        status="created",
        duration_s=time.perf_counter() - start,
    )


def print_banner(year: int, month: int, invoice_dir: Path, dry_run: bool) -> None:
    month_name = dt.date(year, month, 1).strftime("%B")
    mode = "DRY RUN" if dry_run else "LIVE"
    print("=" * 82)
    print(f" 📄 Invoice Generator | {month_name} {year} | Mode: {mode}")
    print(f" 📁 Folder: {invoice_dir}")
    print("=" * 82)


def print_result(index: int, total: int, result: ProcessResult) -> None:
    icon = "✅" if result.status in {"created", "planned"} else "❌"
    print(f"[{index}/{total}] {icon} {result.source_name}")
    if result.status == "planned":
        print(f"    • Plan: {result.docx_name} + {result.pdf_name}")
    else:
        invoice_text = result.invoice_number if result.invoice_number else "(unchanged)"
        print(
            f"    • Output: {result.docx_name} + {result.pdf_name} | "
            f"Invoice#: {invoice_text} | Total: {format_money(result.subtotal)} | "
            f"{result.duration_s:.2f}s"
        )


def maybe_pause(pause_on_exit: bool) -> None:
    if not pause_on_exit:
        return
    try:
        input("\nPress Enter to close...")
    except EOFError:
        pass


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate monthly invoice DOCX and PDF files from February templates."
    )
    parser.add_argument(
        "--invoice-dir",
        type=Path,
        default=DEFAULT_INVOICE_DIR,
        help="Directory containing source invoice DOCX files.",
    )
    parser.add_argument(
        "--year",
        type=int,
        default=dt.date.today().year,
        help="Target year (defaults to current year).",
    )
    parser.add_argument(
        "--month",
        type=int,
        default=dt.date.today().month,
        help="Target month number 1-12 (defaults to current month).",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview output names without writing files.",
    )
    parser.add_argument(
        "--pause-on-exit",
        action="store_true",
        help="Pause at end so the window does not close immediately.",
    )

    args = parser.parse_args()

    if not 1 <= args.month <= 12:
        raise SystemExit("Month must be between 1 and 12")

    run_start = time.perf_counter()
    invoice_files = resolve_invoice_files(args.invoice_dir)
    print_banner(args.year, args.month, args.invoice_dir, args.dry_run)

    completed = 0
    failures = 0

    for index, source_doc in enumerate(invoice_files, start=1):
        try:
            result = process_invoice(source_doc, args.year, args.month, dry_run=args.dry_run)
            print_result(index, len(invoice_files), result)
            completed += 1
        except Exception as exc:  # pragma: no cover - runtime env specific
            failures += 1
            print(f"[{index}/{len(invoice_files)}] ❌ {source_doc.name}")
            print(f"    • Error: {exc}")

    elapsed = time.perf_counter() - run_start
    print("-" * 82)
    print(
        f"Done: {completed} succeeded, {failures} failed, total {len(invoice_files)} invoice(s) "
        f"in {elapsed:.2f}s"
    )

    if failures:
        maybe_pause(args.pause_on_exit)
        raise SystemExit(1)

    maybe_pause(args.pause_on_exit)


if __name__ == "__main__":
    try:
        main()
    except Exception as uncaught_error:
        print(f"\n❌ Fatal error: {uncaught_error}")
        # Keep output visible when launched by double-click in Windows.
        if "--pause-on-exit" in sys.argv:
            try:
                input("Press Enter to close...")
            except EOFError:
                pass
        raise
